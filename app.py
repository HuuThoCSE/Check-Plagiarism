import os
import datetime
from flask import Flask, request, render_template, flash, redirect, url_for
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PyPDF2 import PdfWriter, PdfReader
from requests.exceptions import ConnectionError, ReadTimeout, ConnectTimeout, HTTPError
from googlesearch import search
import logging
import asyncio
import aiohttp
from colorlog import ColoredFormatter
import re
import time
import random
from bs4 import BeautifulSoup
import nltk
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import io
import socket

app = Flask(__name__)
app.secret_key = 'your_secret_key'

# Mã ANSI cho màu trắng
WHITE = '\033[37m'
RESET = '\033[0m'

# Thiết lập logging với colorlog
formatter = ColoredFormatter(
    "%(log_color)s%(asctime)s - %(levelname)s - %(message)s",
    datefmt='%Y-%m-%d %H:%M:%S',
    log_colors={
        'DEBUG': 'cyan',
        'INFO': 'green',
        'WARNING': 'yellow',
        'ERROR': 'red',
        'CRITICAL': 'bold_red',
    }
)

handler = logging.StreamHandler()
handler.setFormatter(formatter)

logging.basicConfig(level=logging.INFO, handlers=[handler])

# Hàm làm sạch văn bản
def clean_text(text):
    text = re.sub(r'\.{2,}', ' ', text)  # Loại bỏ các ký tự "." liên tiếp
    text = re.sub(r'[^\w\s]', '', text)  # Loại bỏ các ký tự không phải là chữ cái hoặc số
    return text.strip()

# Hàm trích xuất nội dung từ file DOCX
def read_docx(file_path):
    logging.info(f"Đang đọc file DOCX: {file_path}")
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        text = clean_text(text)  # Làm sạch văn bản
        if is_valid_paragraph(text):
            full_text.append(text)
    return doc, "\n".join(full_text)

# Hàm trích xuất nội dung từ file PDF
def read_pdf(file_path):
    logging.info(f"Đang đọc file PDF: {file_path}")
    text = ""
    with open(file_path, 'rb') as file:
        reader = PdfReader(file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                for line in page_text.splitlines():
                    line = line.strip()
                    line = clean_text(line)  # Làm sạch văn bản
                    if is_valid_paragraph(line):
                        text += line + "\n"
    return reader, text

# Hàm xác định một đoạn văn có hợp lệ hay không
def is_valid_paragraph(text):
    if len(text) < 50:  # Đoạn quá ngắn
        return False
    if len(re.findall(r'\.{2,}', text)) > 0:  # Có nhiều dấu chấm liên tiếp
        return False
    return True

# Hàm chia nhỏ nội dung thành các đoạn có độ dài tối đa
def split_text_into_chunks(text, max_length=500):
    logging.info("Đang chia nhỏ nội dung thành các đoạn.")
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0

    for word in words:
        if current_length + len(word) + 1 > max_length:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_length = len(word) + 1
        else:
            current_chunk.append(word)
            current_length += len(word) + 1

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks

# Hàm đánh dấu các đoạn văn bị nghi ngờ đạo văn
def mark_plagiarism_chunks(chunks, similarities, sources, threshold=0.8):
    logging.info("Đang đánh dấu các đoạn bị nghi ngờ đạo văn.")
    marked_chunks = []

    for i, (chunk, similarity, source) in enumerate(zip(chunks, similarities, sources)):
        if similarity >= threshold:
            reason = f"Đoạn {i + 1} bị đánh dấu đạo văn với mức độ giống nhau {similarity * 100:.2f}%."
            marked_chunk = f"// Đạo văn: {reason}\n{chunk}\n// Lý do: Đoạn văn này có mức độ giống nhau vượt quá ngưỡng cho phép.\n// Nguồn: {source}\n\n"
            marked_chunks.append(marked_chunk)
            logging.warning(reason)
        else:
            marked_chunks.append(chunk + "\n")

    return "\n".join(marked_chunks)

# Hàm tìm kiếm từ khóa trên Google và trích xuất nội dung từ kết quả
async def fetch_content(session, url):
    try:
        logging.info(f"Đang truy cập: {url}")
        async with session.get(url, timeout=15, ssl=False) as response:  # Tăng thời gian chờ lên 15 giây
            if response.status == 200:
                raw_text = await response.read()
                
                # Thử giải mã với UTF-8 và xử lý các lỗi
                try:
                    text = raw_text.decode('utf-8')
                except UnicodeDecodeError:
                    logging.warning(f"Mã hóa UTF-8 không hợp lệ, thử dùng mã hóa khác cho {url}")
                    text = raw_text.decode('latin1')  # Dùng mã hóa latin1 như một lựa chọn thay thế
                
                soup = BeautifulSoup(text, 'html.parser')
                paragraphs = soup.find_all('p')
                content = ' '.join([para.get_text() for para in paragraphs])
                return content
            else:
                logging.error(f"Lỗi khi truy cập {url}: HTTP {response.status}")
                return ""
    except ConnectionResetError as e:
        logging.error(f"Kết nối bị đóng đột ngột bởi máy chủ {url}: {e}")
        # Thử lại sau khi chờ một khoảng thời gian ngẫu nhiên
        await asyncio.sleep(random.uniform(1, 5))
        return await fetch_content(session, url)
    except Exception as e:
        logging.error(f"Lỗi khi truy cập {url}: {e}")
        return ""

def is_valid_search_query(query):
    if query.count('.') / len(query) > 0.2:  # Nếu hơn 20% nội dung là dấu chấm
        return False
    if len(query.split()) < 3:  # Nếu query có ít hơn 3 từ
        return False
    return True

# Hàm kiểm tra kết nối mạng
def is_connected():
    try:
        # Thử kết nối đến một trang web có kết nối nhanh (như Google)
        socket.create_connection(("www.google.com", 80))
        return True
    except OSError:
        return False

# Cập nhật hàm search_google_async để trả về cả nguồn và nội dung
async def search_google_async(query):
    if not is_valid_search_query(query):
        logging.warning("Query không hợp lệ, bỏ qua tìm kiếm trên Google.")
        return ""

    try:
        logging.info(f"Tìm kiếm trên Google: {query}")
        search_results = search(query, num_results=3)
        content = ""
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=30)) as session:
            tasks = []
            for result in search_results:
                # Nếu search_results trả về quá nhiều giá trị, bạn có thể điều chỉnh để chỉ lấy URL cần thiết
                url = result  # hoặc unpack đúng số lượng giá trị như url, title = result
                tasks.append(fetch_content(session, url))
            contents = await asyncio.gather(*tasks)
            content = ' '.join(contents)
        
        time.sleep(random.uniform(2, 5))
        return content

    except (ReadTimeout, ConnectionError) as e:
        logging.warning(f"Google gặp lỗi kết nối hoặc thời gian chờ: {str(e)}. Chuyển sang tìm kiếm trên Bing...")
        return await search_bing(query)
    except HTTPError as e:
        if e.response.status_code == 429:
            logging.warning("Google trả về lỗi 429. Chuyển sang Bing...")
            return await search_bing(query)
        logging.error(f"Lỗi HTTP không mong muốn: {e}")
        return ""

async def search_bing(query):
    try:
        logging.info(f"Tìm kiếm trên Bing: {query}")
        bing_url = f"https://www.bing.com/search?q={query}"
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=30)) as session:
            bing_response = await fetch_content(session, bing_url)
            if bing_response:
                return bing_response
        logging.warning("Không tìm thấy nội dung phù hợp từ Bing. Chuyển sang Cốc Cốc...")
        return await search_coccoc(query)
    except Exception as e:
        logging.error(f"Lỗi khi tìm kiếm trên Bing: {e}")
        return await search_coccoc(query)

async def search_coccoc(query):
    try:
        logging.info(f"Tìm kiếm trên Cốc Cốc: {WHITE}{query}{RESET}")
        coccoc_url = f"https://coccoc.com/search?query={query}"
        async with aiohttp.ClientSession() as session:
            coccoc_response = await fetch_content(session, coccoc_url)
            if coccoc_response:
                return coccoc_response
        logging.warning("Không tìm thấy nội dung phù hợp từ Cốc Cốc. Chuyển sang Yahoo...")
        return await search_yahoo(query)
    except Exception as e:
        logging.error(f"Lỗi khi tìm kiếm trên Cốc Cốc: {e}")
        return await search_yahoo(query)

async def search_yahoo(query):
    try:
        logging.info(f"Tìm kiếm trên Yahoo: {WHITE}{query}{RESET}")
        yahoo_url = f"https://search.yahoo.com/search?p={query}"
        async with aiohttp.ClientSession() as session:
            yahoo_response = await fetch_content(session, yahoo_url)
            return yahoo_response
    except Exception as e:
        logging.error(f"Lỗi khi tìm kiếm trên Yahoo: {e}")
        return ""

# Hàm tính toán mức độ giống nhau
def calculate_similarity(text1, text2):
    logging.info("Đang tính toán mức độ giống nhau.")
    if not text1.strip() or not text2.strip():
        logging.warning("Một hoặc cả hai đoạn văn bản đều rỗng hoặc chỉ chứa từ dừng.")
        return 0.0
    
    vectorizer = TfidfVectorizer()
    try:
        tfidf_matrix = vectorizer.fit_transform([text1, text2])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
        return min(similarity[0][0], 1.0)  # Giới hạn mức độ giống nhau tối đa là 100%
    except ValueError as e:
        logging.error(f"Lỗi khi tính toán mức độ giống nhau: {e}")
        return 0.0  # Trả về 0 nếu có lỗi xảy ra

def add_mark_to_page(page, chunk, similarity):
    # Tạo một trang mới với dấu đánh dấu
    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    
    # Định vị trí để đánh dấu (tạm thời là ví trí cố định)
    can.drawString(100, 100, f"[Đạo văn] {chunk} - Mức độ giống nhau: {similarity * 100:.2f}%")
    can.save()

    packet.seek(0)
    new_pdf = PdfReader(packet)
    watermark_page = new_pdf.pages[0]

    # Kết hợp watermark_page với page gốc
    page.merge_page(watermark_page)
    return page

# Hàm đánh dấu các đoạn văn bị nghi ngờ đạo văn trong file PDF
def mark_plagiarism_in_pdf(file_path, chunks, similarities, output_path):
    reader = PdfReader(file_path)
    writer = PdfWriter()
    threshold = 0.8

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            page_chunks = split_text_into_chunks(text)
            for chunk, similarity in zip(page_chunks, similarities):
                if similarity >= threshold:
                    # Thêm dấu vào đúng vị trí của đoạn văn
                    page = add_mark_to_page(page, chunk, similarity)
        writer.add_page(page)

    with open(output_path, "wb") as output_pdf:
        writer.write(output_pdf)

# Hàm đánh dấu các đoạn văn bị nghi ngờ đạo văn trong file DOCX với highlight
def mark_plagiarism_in_docx(doc, chunks, similarities, sources):
    threshold = 0.8
    for para in doc.paragraphs:
        for i, (chunk, similarity, source) in enumerate(zip(chunks, similarities, sources)):
            if similarity >= threshold and chunk in para.text:
                run = para.add_run(f"[Đạo văn] {chunk}")
                run.font.highlight_color = docx.shared.WD_COLOR_INDEX.YELLOW
                para.add_run(f" (Nguồn: {source})")
    return doc

# Route chính để tải lên file, tìm kiếm trên Google và kiểm tra đạo văn
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        logging.info("Nhận yêu cầu POST.")
        if 'file' not in request.files:
            flash('Không tìm thấy tệp được tải lên!')
            return redirect(request.url)
        
        file = request.files['file']
        if file.filename == '':
            flash('Không có tệp nào được chọn!')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            logging.info(f"File hợp lệ được tải lên: {file.filename}")
            file_path = os.path.join('uploads', file.filename)
            file.save(file_path)

            # Trích xuất nội dung từ file DOCX hoặc PDF
            if file.filename.endswith('.docx'):
                doc, file_content = read_docx(file_path)
            elif file.filename.endswith('.pdf'):
                reader, file_content = read_pdf(file_path)
            else:
                flash('Định dạng tệp không được hỗ trợ!')
                return redirect(request.url)

            # Chia nhỏ nội dung thành các đoạn
            chunks = split_text_into_chunks(file_content)

            overall_similarity = 0
            similarities = []
            sources = []
            for i, chunk in enumerate(chunks):
                # Tìm kiếm trên Google và trích xuất nội dung
                try:
                    search_content, search_sources = asyncio.run(search_google_async(chunk))
                except ValueError as e:
                    logging.error(f"Error unpacking search results: {e}")
                    search_content, search_sources = "", []

                # Kiểm tra nếu search_content rỗng
                if not search_content.strip():
                    logging.warning(f"Không tìm thấy nội dung phù hợp cho đoạn {i+1}.")
                    # Nếu không tìm thấy nội dung, chia nhỏ đoạn văn bản
                    smaller_chunks = split_text_into_chunks(chunk, max_length=250)
                    for smaller_chunk in smaller_chunks:
                        try:
                            search_content, search_sources = asyncio.run(search_google_async(smaller_chunk))
                        except ValueError as e:
                            logging.error(f"Error unpacking search results: {e}")
                            search_content, search_sources = "", []
                        
                        similarity = calculate_similarity(smaller_chunk, search_content)
                        overall_similarity += similarity
                        similarities.append(similarity)
                        sources.extend(search_sources)
                    continue

                # Kiểm tra nội dung trùng lặp giữa đoạn văn bản và kết quả tìm kiếm
                similarity = calculate_similarity(chunk, search_content)
                overall_similarity += similarity
                similarities.append(similarity)
                sources.extend(search_sources)
                
                # Log thông tin tiến trình
                percentage_completed = ((i + 1) / len(chunks)) * 100
                logging.info(f'Đoạn {i + 1}/{len(chunks)}: Mức độ giống nhau: {similarity * 100:.2f}%, Hoàn thành: {percentage_completed:.2f}%')

            # Đánh dấu các đoạn văn bị nghi ngờ đạo văn
            marked_content = mark_plagiarism_chunks(chunks, similarities, sources)

            # Lưu file đã được đánh dấu với thời gian hiện tại
            current_time = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
            save_dir = "checked_files"
            os.makedirs(save_dir, exist_ok=True)

            if file.filename.endswith('.docx'):
                marked_doc = mark_plagiarism_in_docx(doc, chunks, similarities, sources)
                output_path = os.path.join(save_dir, f"checked_{current_time}.docx")
                marked_doc.save(output_path)
            elif file.filename.endswith('.pdf'):
                output_path = os.path.join(save_dir, f"checked_{current_time}.pdf")
                mark_plagiarism_in_pdf(file_path, chunks, similarities, output_path)

            # Ghi kết quả vào tệp văn bản
            text_output_path = os.path.join(save_dir, f"marked_{current_time}.txt")
            with open(text_output_path, 'w', encoding='utf-8') as f:
                f.write(marked_content)

            # Tính mức độ giống nhau trung bình
            if similarities:
                average_similarity = (overall_similarity / len(similarities)) * 100
            else:
                average_similarity = 0

            flash(f'Mức độ giống nhau trung bình giữa tệp và nội dung tìm thấy trên Google: {average_similarity:.2f}%')
            logging.info(f'Tổng số đoạn: {len(chunks)}, Mức độ giống nhau trung bình: {average_similarity:.2f}%')

            return render_template('result.html', similarity=average_similarity)

    return render_template('index.html')


# Hàm kiểm tra định dạng file được phép
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'docx', 'pdf'}

if __name__ == '__main__':
    os.makedirs('uploads', exist_ok=True)
    app.run(debug=True)
